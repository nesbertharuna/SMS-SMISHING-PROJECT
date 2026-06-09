"""Generate Word-ready thesis appendices document."""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


REPO = Path(__file__).resolve().parents[1]
OUT_DOCX = REPO / "docs" / "THESIS_APPENDICES.docx"

DATASET_ROWS = [
    (1, "benign", "SMS Service", "Econet: You have successfully purchased a WhatsApp bundle."),
    (2, "smishing", "SMS Service", "Final warning: SIM registration incomplete. Update details immediately."),
    (3, "smishing", "SMS Service", "Econet Alert: Your line will be disconnected in 24 hours. Verify now at http://econet-security-update.com"),
    (4, "benign", "Rewards", "Happy birthday! Wishing you a wonderful year ahead."),
    (5, "benign", "Lifestyle", "Class timetable updated. Please check the university portal."),
    (6, "benign", "Bank/Finance", "Your EcoCash transaction of USD 5.00 was completed successfully."),
    (7, "smishing", "Bank/Finance", "You are selected for a loan approval. Send national ID and PIN."),
    (8, "benign", "Lifestyle", "Reminder: Your appointment at Bindura Hospital is tomorrow at 10AM."),
    (9, "benign", "Lifestyle", "Your Mukuru transfer is ready for collection at the nearest branch."),
    (10, "smishing", "Accounts", "Customer Service: Your account has unusual login attempts. Confirm password urgently."),
    (11, "smishing", "Accounts", "WhatsApp account compromised. Click www.verify-whatsapp-login.com now."),
    (12, "benign", "SMS Service", "NetOne: Recharge successful. You received 500MB bonus data."),
    (13, "benign", "Lifestyle", "Your package is ready for pickup at the local post office."),
    (14, "benign", "Lifestyle", "Church service starts at 9AM this Sunday in Bindura."),
    (15, "smishing", "Bank/Finance", "ZIMRA Tax Refund: Submit your banking details to receive payment."),
    (16, "smishing", "Bank/Finance", "Congratulations! You won USD500 EcoCash reward. Claim now."),
    (17, "smishing", "Accounts", "Your Mukuru account was flagged. Verify identity now to avoid suspension."),
    (18, "smishing", "Bank/Finance", "EcoCash: Your wallet has been suspended. Confirm PIN immediately."),
    (19, "smishing", "Rewards", "NetOne Promotion! Claim your free airtime reward here: bit.ly/free-netone"),
    (20, "benign", "Lifestyle", "ZESA: Your token purchase was successful."),
]

CODE_FILES = [
    ("A.1", "SMS Text Preprocessing", REPO / "ml-backend/preprocessing/clean_text.py"),
    ("A.2", "TF-IDF Feature Extraction and Logistic Regression Pipeline", REPO / "ml-backend/feature-extraction/tfidf_features.py"),
    ("A.3", "Model Training Script", REPO / "experiments/train.py"),
    ("A.4", "ML Model Evaluation Script", REPO / "experiments/evaluate_ml.py"),
    ("A.5", "Rule-Based Baseline Evaluation Script", REPO / "experiments/evaluate_rules.py"),
    ("A.6", "Rule-Based Smishing Detector", REPO / "ml-backend/rule-engine/rules.py"),
    ("A.7", "Logistic Regression Explanation Module", REPO / "ml-backend/models/explain_lr_tfidf.py"),
    ("A.8", "Dataset Loader", REPO / "scripts/load_dataset1.py"),
    ("A.9", "Stratified Data Splitting", REPO / "scripts/make_splits.py"),
]


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def add_code_block(doc: Document, code: str, filename: str) -> None:
    add_body(doc, f"Source file: {filename}")
    for line in code.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "D9E2F3") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def add_confusion_matrix_section(
    doc: Document,
    title: str,
    tn: int,
    fp: int,
    fn: int,
    tp: int,
    n_test: int,
    extra_rows: list[list[str]] | None = None,
) -> None:
    add_heading(doc, title, level=2)
    add_body(doc, f"Test set size (n): {n_test}")
    add_body(doc, "Label convention: 0 = benign (ham), 1 = smishing (spam). Rows = actual, columns = predicted.")
    add_table(
        doc,
        ["", "Predicted 0 (benign)", "Predicted 1 (smishing)"],
        [
            ["Actual 0 (benign)", str(tn), str(fp)],
            ["Actual 1 (smishing)", str(fn), str(tp)],
        ],
    )
    add_table(
        doc,
        ["Cell", "Name", "Count", "Meaning"],
        [
            ["Top-left", "TN (True Negative)", str(tn), "Benign correctly classified as benign"],
            ["Top-right", "FP (False Positive)", str(fp), "Benign incorrectly classified as smishing"],
            ["Bottom-left", "FN (False Negative)", str(fn), "Smishing incorrectly classified as benign"],
            ["Bottom-right", "TP (True Positive)", str(tp), "Smishing correctly classified as smishing"],
        ],
    )
    if extra_rows:
        add_table(doc, ["Metric", "Value"], extra_rows)


def main() -> None:
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("APPENDICES AND SUPPLEMENTARY MATERIALS")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("SMS Smishing Detection System — TF-IDF + Logistic Regression")
    sub_run.font.size = Pt(12)

    doc.add_paragraph()
    add_body(
        doc,
        "This document contains supplementary material for the smishing detection system, including "
        "complete source code listings, dataset samples, hyperparameter configurations, and confusion matrices.",
    )

    # Table of contents style list
    add_heading(doc, "Contents", level=1)
    for item in [
        "Appendix A — Complete Code Listings",
        "Appendix B — Dataset Samples and Statistics",
        "Appendix C — Hyperparameter Configurations",
        "Appendix D — Complete Confusion Matrices",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # APPENDIX A
    add_heading(doc, "Appendix A — Complete Code Listings", level=1)
    add_body(
        doc,
        "The following listings reproduce the core Python modules used to preprocess SMS text, "
        "extract TF-IDF features, train the Logistic Regression classifier, evaluate performance, "
        "and run the rule-based baseline. File paths are relative to the project repository root.",
    )

    for section_id, title_text, path in CODE_FILES:
        add_heading(doc, f"{section_id}  {title_text}", level=2)
        code = path.read_text(encoding="utf-8")
        add_code_block(doc, code, str(path.relative_to(REPO)).replace("\\", "/"))
        doc.add_page_break()

    # APPENDIX B
    add_heading(doc, "Appendix B — Dataset Samples and Statistics", level=1)

    add_heading(doc, "B.1  Dataset Schema", level=2)
    add_body(doc, "Raw dataset file: ml-backend/dataset/Dataset1.csv")
    add_table(
        doc,
        ["Column", "Type", "Description"],
        [
            ["label", "string", "Class label: benign or smishing"],
            ["message", "string", "Raw SMS message text"],
            ["feature_notes", "string", "Category tag for documentation (not used in model training)"],
        ],
    )

    add_heading(doc, "B.2  Dataset Statistics", level=2)
    raw = pd.read_csv(REPO / "ml-backend/dataset/Dataset1.csv")
    add_table(
        doc,
        ["Statistic", "Value"],
        [
            ["Raw CSV rows", str(len(raw))],
            ["Raw benign rows", str((raw["label"] == "benign").sum())],
            ["Raw smishing rows", str((raw["label"] == "smishing").sum())],
            ["Unique messages after deduplication", "20"],
            ["Benign unique messages", "10"],
            ["Smishing unique messages", "10"],
            ["Train / validation / test split", "14 / 3 / 3 (70% / 15% / 15%)"],
            ["Random seed", "42"],
            ["Split method", "Stratified (sklearn train_test_split, stratify=label)"],
        ],
    )
    add_body(
        doc,
        "Note: The loader script (load_dataset1.py) deduplicates messages by text. If the same message "
        "appears with conflicting labels, the smishing label is retained. This reduces 600 raw rows to "
        "20 unique modelling examples.",
    )

    add_heading(doc, "B.3  Complete Unique Message Catalogue", level=2)
    add_table(
        doc,
        ["No.", "Label", "Category", "Message"],
        [[str(a), b, c, d] for a, b, c, d in DATASET_ROWS],
    )

    add_heading(doc, "B.4  Test Set Holdout Messages", level=2)
    add_table(
        doc,
        ["Label", "Numeric label", "Message"],
        [
            ["benign", "0", "Your EcoCash transaction of USD 5.00 was completed successfully."],
            ["benign", "0", "NetOne: Recharge successful. You received 500MB bonus data."],
            ["smishing", "1", "EcoCash: Your wallet has been suspended. Confirm PIN immediately."],
        ],
    )

    doc.add_page_break()

    # APPENDIX C
    add_heading(doc, "Appendix C — Hyperparameter Configurations", level=1)

    add_heading(doc, "C.1  Data Splitting Parameters", level=2)
    add_table(
        doc,
        ["Parameter", "Value", "Description"],
        [
            ["train_ratio", "0.70", "Proportion of data for training"],
            ["val_ratio", "0.15", "Proportion for validation"],
            ["test_ratio", "0.15", "Proportion for held-out test"],
            ["random_state", "42", "Reproducible random seed"],
            ["stratify", "label", "Preserves class proportions in each split"],
        ],
    )

    add_heading(doc, "C.2  Text Preprocessing Parameters", level=2)
    add_table(
        doc,
        ["Parameter", "Value"],
        [
            ["Lowercasing", "Yes"],
            ["Punctuation removal", "Yes (regex [^\\w\\s]+)"],
            ["Whitespace normalization", "Yes"],
            ["Stopword removal", "Yes (built-in English list, ~130 words)"],
            ["keep_numbers", "True"],
            ["Tokenization", "Space-split after preprocessing"],
        ],
    )

    add_heading(doc, "C.3  TF-IDF Vectorizer Parameters (TfidfConfig)", level=2)
    add_table(
        doc,
        ["Parameter", "Value", "Description"],
        [
            ["max_features", "20000", "Maximum vocabulary size"],
            ["ngram_range", "(1, 2)", "Unigrams and bigrams"],
            ["min_df", "2", "Ignore terms appearing fewer than 2 times"],
            ["max_df", "0.95", "Ignore terms in more than 95% of documents"],
            ["sublinear_tf", "True", "Apply sublinear term frequency scaling"],
            ["tokenizer", "str.split", "Space-separated tokens from preprocessor"],
            ["lowercase", "False", "Already lowercased in preprocessing step"],
        ],
    )

    add_heading(doc, "C.4  Logistic Regression Classifier Parameters (Production Model)", level=2)
    add_table(
        doc,
        ["Parameter", "Value", "Description"],
        [
            ["Algorithm", "sklearn.linear_model.LogisticRegression", "Binary classifier"],
            ["max_iter", "2000", "Maximum optimisation iterations"],
            ["class_weight", "balanced", "Adjusts for class imbalance"],
            ["n_jobs", "None", "Single-threaded"],
            ["Saved artifact", "ml-backend/artifacts/pipeline_lr_tfidf.joblib", "Serialised pipeline"],
        ],
    )

    add_heading(doc, "C.5  Rule-Based Baseline Parameters (RuleConfig)", level=2)
    add_table(
        doc,
        ["Parameter", "Value", "Description"],
        [
            ["threshold", "6", "Score at or above this value → smishing"],
            ["keyword_weight", "2", "Points per matched risky keyword"],
            ["url_weight", "4", "Points if URL detected"],
            ["phone_weight", "2", "Points if phone number detected"],
            ["money_weight", "1", "Points if monetary amount detected"],
            ["Phrase bonus examples", "verify your account (+3), account suspended (+3)", "Extra phrase weights"],
            ["Urgency cues", "≥2 exclamation marks (+1); ≥2 ALL-CAPS tokens (+1)", "Punctuation/caps heuristics"],
        ],
    )

    add_body(doc, "Monitored keywords (24 terms):")
    add_body(
        doc,
        "urgent, verify, account, password, login, confirm, suspended, locked, security, fraud, "
        "unauthorized, reset, update, payment, billing, refund, claim, winner, prize, free, offer, "
        "limited, act, now",
    )

    add_heading(doc, "C.6  Software Dependencies", level=2)
    add_table(
        doc,
        ["Package", "Version constraint"],
        [
            ["numpy", "≥2.4, <3.0"],
            ["pandas", "≥3.0, <4.0"],
            ["scikit-learn", "≥1.8, <2.0"],
            ["joblib", "≥1.5, <2.0"],
            ["scipy", "≥1.17, <2.0"],
        ],
    )

    doc.add_page_break()

    # APPENDIX D
    add_heading(doc, "Appendix D — Complete Confusion Matrices", level=1)
    add_body(
        doc,
        "All matrices below use binary labels: 0 = benign, 1 = smishing. "
        "Evaluation run date: 2026-05-30. Model: TF-IDF + Logistic Regression (class_weight=balanced). "
        "Random seed: 42.",
    )

    add_heading(doc, "D.1  ML Model — Validation Set", level=2)
    add_confusion_matrix_section(
        doc,
        "Validation set confusion matrix (n = 3)",
        tn=1,
        fp=0,
        fn=1,
        tp=1,
        n_test=3,
        extra_rows=[
            ["Accuracy", "66.7%"],
            ["Precision (class 1)", "100.0%"],
            ["Recall (class 1)", "50.0%"],
            ["F1-score (class 1)", "66.7%"],
        ],
    )

    ml_metrics = json.loads((REPO / "reports/metrics_ml_test.json").read_text(encoding="utf-8"))
    cm = ml_metrics["confusion_matrix_named"]
    m = ml_metrics["metrics_pos1"]
    cr = ml_metrics["classification_report"]

    add_heading(doc, "D.2  ML Model — Test Set (Held-Out)", level=2)
    add_confusion_matrix_section(
        doc,
        "Test set confusion matrix (n = 3)",
        tn=cm["tn"],
        fp=cm["fp"],
        fn=cm["fn"],
        tp=cm["tp"],
        n_test=ml_metrics["n_test"],
        extra_rows=[
            ["Accuracy", f"{cr['accuracy'] * 100:.1f}%"],
            ["Precision (class 0 — benign)", f"{cr['0']['precision'] * 100:.1f}%"],
            ["Recall (class 0 — benign)", f"{cr['0']['recall'] * 100:.1f}%"],
            ["F1-score (class 0 — benign)", f"{cr['0']['f1-score'] * 100:.1f}%"],
            ["Precision (class 1 — smishing)", f"{cr['1']['precision'] * 100:.1f}%"],
            ["Recall (class 1 — smishing)", f"{cr['1']['recall'] * 100:.1f}%"],
            ["F1-score (class 1 — smishing)", f"{cr['1']['f1-score'] * 100:.1f}%"],
        ],
    )

    add_heading(doc, "D.3  Per-Row Test Predictions (ML Model)", level=2)
    pred_df = pd.read_csv(REPO / "reports/pred_ml_test.csv")
    test_df = pd.read_csv(REPO / "data/splits/test.csv")
    rows = []
    for i in range(len(pred_df)):
        label_name = "benign" if pred_df.iloc[i]["y_true"] == 0 else "smishing"
        pred_name = "benign" if pred_df.iloc[i]["y_pred"] == 0 else "smishing"
        correct = "Yes" if pred_df.iloc[i]["y_true"] == pred_df.iloc[i]["y_pred"] else "No"
        msg = test_df.iloc[i]["message"]
        rows.append([str(i + 1), label_name, pred_name, correct, msg])
    add_table(doc, ["Row", "Actual", "Predicted", "Correct?", "Message"], rows)

    rules_metrics = json.loads((REPO / "reports/metrics_rules_test.json").read_text(encoding="utf-8"))
    rcm = rules_metrics["confusion_matrix_labels_[0,1]"]
    rcr = rules_metrics["classification_report"]

    add_heading(doc, "D.4  Rule-Based Baseline — Test Set", level=2)
    add_confusion_matrix_section(
        doc,
        f"Rule-based test confusion matrix (threshold = {rules_metrics['rule_threshold']}, n = {rules_metrics['n_test']})",
        tn=rcm[0][0],
        fp=rcm[0][1],
        fn=rcm[1][0],
        tp=rcm[1][1],
        n_test=rules_metrics["n_test"],
        extra_rows=[
            ["Accuracy", f"{rcr['accuracy'] * 100:.1f}%"],
            ["Precision (class 1 — smishing)", f"{rcr['1']['precision'] * 100:.1f}%"],
            ["Recall (class 1 — smishing)", f"{rcr['1']['recall'] * 100:.1f}%"],
            ["F1-score (class 1 — smishing)", f"{rcr['1']['f1-score'] * 100:.1f}%"],
        ],
    )

    add_heading(doc, "D.5  Full Classification Report — ML Test Set", level=2)
    add_table(
        doc,
        ["Class", "Precision", "Recall", "F1-Score", "Support"],
        [
            ["0 (benign)", f"{cr['0']['precision']:.4f}", f"{cr['0']['recall']:.4f}", f"{cr['0']['f1-score']:.4f}", str(int(cr['0']['support']))],
            ["1 (smishing)", f"{cr['1']['precision']:.4f}", f"{cr['1']['recall']:.4f}", f"{cr['1']['f1-score']:.4f}", str(int(cr['1']['support']))],
            ["Accuracy", f"{cr['accuracy']:.4f}", "—", "—", str(int(cr['weighted avg']['support']))],
            ["Macro avg", f"{cr['macro avg']['precision']:.4f}", f"{cr['macro avg']['recall']:.4f}", f"{cr['macro avg']['f1-score']:.4f}", str(int(cr['macro avg']['support']))],
            ["Weighted avg", f"{cr['weighted avg']['precision']:.4f}", f"{cr['weighted avg']['recall']:.4f}", f"{cr['weighted avg']['f1-score']:.4f}", str(int(cr['weighted avg']['support']))],
        ],
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Saved -> {OUT_DOCX}")


if __name__ == "__main__":
    main()
